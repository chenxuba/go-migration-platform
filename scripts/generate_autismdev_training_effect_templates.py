#!/usr/bin/env python3
"""Generate DOCX and PNG templates for AutismDev training effect tables."""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
from functools import lru_cache
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[1]
ITEM_BANK = REPO_ROOT / "docs" / "autismdev-item-bank-draft.json"
ASSET_DIR = (
    REPO_ROOT
    / "services"
    / "education"
    / "internal"
    / "service"
    / "assets"
    / "autismdev_training_effect"
)
DEFAULT_DOCX_DIR = REPO_ROOT / "output" / "autismdev_training_effect_templates"


DOMAIN_NAMES = {
    "SP": "感知觉",
    "GM": "粗大动作",
    "FM": "精细动作",
    "LC": "语言与沟通",
    "COG": "认知",
    "SOC": "社会交往",
    "ADL": "生活自理",
    "EB": "情绪与行为",
}

DOMAIN_SECTIONS = {
    "SP": "4.1",
    "GM": "4.2",
    "FM": "4.3",
    "LC": "4.4",
    "COG": "4.5",
    "SOC": "4.6",
    "ADL": "4.7",
    "EB": "4.8",
}

PAGES = [
    (43, "SP", 1, 19),
    (44, "SP", 20, 36),
    (45, "SP", 37, 55),
    (46, "GM", 1, 27),
    (47, "GM", 28, 53),
    (48, "GM", 54, 72),
    (49, "FM", 1, 21),
    (50, "FM", 22, 43),
    (51, "FM", 44, 66),
    (52, "LC", 1, 25),
    (53, "LC", 26, 54),
    (54, "LC", 55, 79),
    (55, "COG", 1, 26),
    (56, "COG", 27, 55),
    (57, "SOC", 1, 16),
    (58, "SOC", 17, 35),
    (59, "SOC", 36, 47),
    (60, "ADL", 1, 22),
    (61, "ADL", 23, 46),
    (62, "ADL", 47, 67),
    (63, "EB", 1, 27),
    (64, "EB", 28, 52),
]

COLUMN_WIDTHS_MM = [14, 60, 18, 18, 14, 14, 14, 14, 14, 14]
PNG_PAGE_WIDTH = 1141
PNG_PAGE_HEIGHT = 1600
PNG_TABLE_LEFT = 66
PNG_TABLE_RIGHT = 1075
PNG_TABLE_TOP = 180
PNG_TABLE_BOTTOM = 1468
PNG_HEADER_HEIGHTS = [60, 44, 44]
PNG_LINE_WIDTH = 2
PNG_TITLE_FONT_SIZE = 34
PNG_HEADER_FONT_SIZE = 24
PNG_SUBHEADER_FONT_SIZE = 21
PNG_BODY_FONT_SIZE = 20
PNG_SMALL_BODY_FONT_SIZE = 18
PNG_TEXT_COLOR = (0, 0, 0)

SONG_FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Songti.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
]
HEI_FONT_CANDIDATES = [
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
]


def set_cell_border(cell, size: int = 8, color: str = "000000") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 45, start: int = 70, bottom: int = 45, end: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_height(row, height_mm: float, exact: bool = True) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    height = tr_pr.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        tr_pr.append(height)
    height.set(qn("w:val"), str(int(height_mm * 56.7)))
    height.set(qn("w:hRule"), "exact" if exact else "atLeast")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_width(table, width_mm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_grid(table, widths_mm: list[float]) -> None:
    tbl_grid = table._tbl.tblGrid
    for index, width in enumerate(widths_mm):
        if index >= len(tbl_grid.gridCol_lst):
            grid_col = OxmlElement("w:gridCol")
            tbl_grid.append(grid_col)
        else:
            grid_col = tbl_grid.gridCol_lst[index]
        grid_col.set(qn("w:w"), str(int(width * 56.7)))


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()


def write_cell(cell, text: str, size: float = 10.0, bold: bool = False, center: bool = True) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def setup_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(8)
    section.right_margin = Mm(8)
    for style_name in ("Normal",):
        style = document.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(10)
    return document


def add_title(document: Document, page_no: int, code: str) -> None:
    title = f"{DOMAIN_SECTIONS[code]} 孤独症儿童{DOMAIN_NAMES[code]}训练效果评估表"
    if page_no not in (43, 46, 49, 52, 55, 57, 60, 63):
        title += "（续）"
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(title)
    run.font.name = "SimHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    run.font.size = Pt(16)
    run.bold = True


def add_training_table(document: Document, page_no: int, code: str, first_no: int, last_no: int, item_by_domain_no: dict[int, dict]) -> None:
    item_count = last_no - first_no + 1
    table = document.add_table(rows=3 + item_count, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, sum(COLUMN_WIDTHS_MM))
    set_table_grid(table, COLUMN_WIDTHS_MM)
    for col_idx, width in enumerate(COLUMN_WIDTHS_MM):
        for cell in table.columns[col_idx].cells:
            cell.width = Mm(width)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_shading(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).merge(table.cell(2, 1))
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(0, 9))
    table.cell(1, 2).merge(table.cell(2, 2))
    table.cell(1, 3).merge(table.cell(2, 3))
    table.cell(1, 4).merge(table.cell(1, 6))
    table.cell(1, 7).merge(table.cell(1, 9))

    for row_index in range(3):
        set_repeat_header(table.rows[row_index])
    set_row_height(table.rows[0], 11)
    set_row_height(table.rows[1], 8)
    set_row_height(table.rows[2], 8)

    write_cell(table.cell(0, 0), "代号", size=11, bold=True)
    write_cell(table.cell(0, 1), "项    目", size=11, bold=True)
    write_cell(table.cell(0, 2), "训练项目", size=11, bold=True)
    write_cell(table.cell(0, 4), "训练效果", size=11, bold=True)
    write_cell(table.cell(1, 2), "第一次", size=10)
    write_cell(table.cell(1, 3), "第二次", size=10)
    write_cell(table.cell(1, 4), "第一次", size=10)
    write_cell(table.cell(1, 7), "第二次", size=10)
    for col, label in zip(range(4, 10), ["显效", "有效", "无效", "显效", "有效", "无效"]):
        write_cell(table.cell(2, col), label, size=9.5)

    row_height = body_row_height(item_count)
    for offset, item_no in enumerate(range(first_no, last_no + 1)):
        row = table.rows[3 + offset]
        set_row_height(row, row_height)
        item = item_by_domain_no[item_no]
        write_cell(row.cells[0], str(item_no), size=9.5)
        write_cell(row.cells[1], normalize_title(item["item_title"]), size=item_font_size(item["item_title"], item_count), center=False)
        for col in range(2, 10):
            write_cell(row.cells[col], "", size=9.5)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.keep_with_next = False
    document.add_paragraph()


def body_row_height(item_count: int) -> float:
    if item_count >= 28:
        return 7.55
    if item_count >= 26:
        return 8.05
    if item_count >= 24:
        return 8.7
    if item_count >= 21:
        return 9.8
    if item_count >= 18:
        return 11.0
    if item_count >= 16:
        return 12.8
    return 16.0


def item_font_size(title: str, item_count: int) -> float:
    length = len(title)
    if item_count >= 28:
        return 7.2 if length > 12 else 7.8
    if item_count >= 24:
        return 7.8 if length > 14 else 8.4
    if item_count >= 21:
        return 8.2 if length > 16 else 8.8
    if length > 22:
        return 8.4
    return 9.2


def normalize_title(value: str) -> str:
    return value.replace(",", "，")


def table_column_xs() -> list[float]:
    total = sum(COLUMN_WIDTHS_MM)
    width = PNG_TABLE_RIGHT - PNG_TABLE_LEFT
    xs = [float(PNG_TABLE_LEFT)]
    cursor = float(PNG_TABLE_LEFT)
    for column_width in COLUMN_WIDTHS_MM:
        cursor += width * column_width / total
        xs.append(cursor)
    xs[-1] = float(PNG_TABLE_RIGHT)
    return xs


def table_row_ys(item_count: int) -> list[float]:
    ys = [float(PNG_TABLE_TOP)]
    cursor = float(PNG_TABLE_TOP)
    for height in PNG_HEADER_HEIGHTS:
        cursor += height
        ys.append(cursor)
    body_height = (PNG_TABLE_BOTTOM - cursor) / item_count
    for _ in range(item_count):
        cursor += body_height
        ys.append(cursor)
    ys[-1] = float(PNG_TABLE_BOTTOM)
    return ys


@lru_cache(maxsize=32)
def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = HEI_FONT_CANDIDATES if bold else SONG_FONT_CANDIDATES
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_bbox(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int, int, int]:
    return draw.textbbox((0, 0), text, font=font)


def text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    left, _, right, _ = text_bbox(draw, text, font)
    return right - left


def split_wrapped_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: float) -> list[str]:
    text = text.strip()
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if current and text_width(draw, candidate, font) > max_width:
            lines.append(current)
            current = char
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def draw_text_in_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[float, float, float, float],
    text: str,
    font_size: int,
    bold: bool = False,
    align: str = "center",
    padding_x: int = 8,
    max_lines: int | None = None,
) -> None:
    left, top, right, bottom = box
    font = load_font(font_size, bold)
    lines = split_wrapped_text(draw, text, font, max(1, right - left - padding_x * 2))
    if max_lines is not None and len(lines) > max_lines:
        lines = lines[:max_lines]
    line_heights = []
    for line in lines:
        bbox = text_bbox(draw, line, font)
        line_heights.append(bbox[3] - bbox[1])
    line_gap = max(2, int(font_size * 0.12))
    total_height = sum(line_heights) + line_gap * max(0, len(lines) - 1)
    y = top + (bottom - top - total_height) / 2
    for index, line in enumerate(lines):
        bbox = text_bbox(draw, line, font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        if align == "left":
            x = left + padding_x
        else:
            x = left + (right - left - width) / 2
        draw.text((x, y - bbox[1]), line, font=font, fill=PNG_TEXT_COLOR)
        y += height + line_gap


def draw_center_text(draw: ImageDraw.ImageDraw, center_x: float, center_y: float, text: str, font_size: int, bold: bool = False) -> None:
    font = load_font(font_size, bold)
    bbox = text_bbox(draw, text, font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((center_x - width / 2 - bbox[0], center_y - height / 2 - bbox[1]), text, font=font, fill=PNG_TEXT_COLOR)


def draw_table_line(draw: ImageDraw.ImageDraw, xy: tuple[float, float, float, float]) -> None:
    draw.line(tuple(round(value) for value in xy), fill=PNG_TEXT_COLOR, width=PNG_LINE_WIDTH)


def draw_training_table_png(
    out_png: Path,
    page_no: int,
    code: str,
    first_no: int,
    last_no: int,
    item_by_domain_no: dict[int, dict],
) -> None:
    item_count = last_no - first_no + 1
    image = Image.new("RGB", (PNG_PAGE_WIDTH, PNG_PAGE_HEIGHT), "white")
    draw = ImageDraw.Draw(image)
    title = f"{DOMAIN_SECTIONS[code]} 孤独症儿童{DOMAIN_NAMES[code]}训练效果评估表"
    if page_no not in (43, 46, 49, 52, 55, 57, 60, 63):
        title += "（续）"
    draw_center_text(draw, PNG_PAGE_WIDTH / 2, 150, title, PNG_TITLE_FONT_SIZE, bold=True)

    xs = table_column_xs()
    ys = table_row_ys(item_count)

    # Horizontal table lines, leaving merged header cells visually intact.
    draw_table_line(draw, (xs[0], ys[0], xs[-1], ys[0]))
    draw_table_line(draw, (xs[2], ys[1], xs[-1], ys[1]))
    draw_table_line(draw, (xs[4], ys[2], xs[-1], ys[2]))
    for y in ys[3:]:
        draw_table_line(draw, (xs[0], y, xs[-1], y))

    # Vertical lines with merged header cells matching the manual table.
    draw_table_line(draw, (xs[0], ys[0], xs[0], ys[-1]))
    draw_table_line(draw, (xs[1], ys[0], xs[1], ys[-1]))
    draw_table_line(draw, (xs[2], ys[0], xs[2], ys[-1]))
    draw_table_line(draw, (xs[3], ys[1], xs[3], ys[-1]))
    draw_table_line(draw, (xs[4], ys[0], xs[4], ys[-1]))
    draw_table_line(draw, (xs[5], ys[2], xs[5], ys[-1]))
    draw_table_line(draw, (xs[6], ys[2], xs[6], ys[-1]))
    draw_table_line(draw, (xs[7], ys[1], xs[7], ys[-1]))
    draw_table_line(draw, (xs[8], ys[2], xs[8], ys[-1]))
    draw_table_line(draw, (xs[9], ys[2], xs[9], ys[-1]))
    draw_table_line(draw, (xs[10], ys[0], xs[10], ys[-1]))

    draw_text_in_box(draw, (xs[0], ys[0], xs[1], ys[3]), "代号", PNG_HEADER_FONT_SIZE, bold=True)
    draw_text_in_box(draw, (xs[1], ys[0], xs[2], ys[3]), "项    目", PNG_HEADER_FONT_SIZE, bold=True)
    draw_text_in_box(draw, (xs[2], ys[0], xs[4], ys[1]), "训练项目", PNG_HEADER_FONT_SIZE, bold=True)
    draw_text_in_box(draw, (xs[4], ys[0], xs[10], ys[1]), "训练效果", PNG_HEADER_FONT_SIZE, bold=True)
    draw_text_in_box(draw, (xs[2], ys[1], xs[3], ys[3]), "第一次", PNG_SUBHEADER_FONT_SIZE)
    draw_text_in_box(draw, (xs[3], ys[1], xs[4], ys[3]), "第二次", PNG_SUBHEADER_FONT_SIZE)
    draw_text_in_box(draw, (xs[4], ys[1], xs[7], ys[2]), "第一次", PNG_SUBHEADER_FONT_SIZE)
    draw_text_in_box(draw, (xs[7], ys[1], xs[10], ys[2]), "第二次", PNG_SUBHEADER_FONT_SIZE)
    for col, label in zip(range(4, 10), ["显效", "有效", "无效", "显效", "有效", "无效"]):
        draw_text_in_box(draw, (xs[col], ys[2], xs[col + 1], ys[3]), label, PNG_SUBHEADER_FONT_SIZE)

    for offset, item_no in enumerate(range(first_no, last_no + 1)):
        top = ys[3 + offset]
        bottom = ys[4 + offset]
        item = item_by_domain_no[item_no]
        title_text = normalize_title(item["item_title"])
        font_size = PNG_BODY_FONT_SIZE if item_count <= 24 else PNG_SMALL_BODY_FONT_SIZE
        draw_text_in_box(draw, (xs[0], top, xs[1], bottom), str(item_no), PNG_BODY_FONT_SIZE)
        draw_text_in_box(draw, (xs[1], top, xs[2], bottom), title_text, font_size, align="left", padding_x=10, max_lines=2)

    image.save(out_png)


def build_page(document: Document, page: tuple[int, str, int, int], items_by_code: dict[str, dict[int, dict]]) -> None:
    page_no, code, first_no, last_no = page
    add_title(document, page_no, code)
    add_training_table(document, page_no, code, first_no, last_no, items_by_code[code])


def load_items() -> dict[str, dict[int, dict]]:
    items = json.loads(ITEM_BANK.read_text(encoding="utf-8"))
    out: dict[str, dict[int, dict]] = {}
    for item in items:
        out.setdefault(item["domain_code"], {})[item["domain_item_no"]] = item
    return out


def render_docx_thumbnail(docx_path: Path, out_png: Path) -> None:
    temp_dir = out_png.parent
    generated = temp_dir / f"{docx_path.name}.png"
    if generated.exists():
        generated.unlink()
    result = subprocess.run(
        ["qlmanage", "-t", "-s", "1600", "-o", str(temp_dir), str(docx_path)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    if result.returncode != 0 or not generated.exists():
        raise RuntimeError(f"qlmanage failed for {docx_path}: {result.stdout}{result.stderr}")
    if out_png.exists():
        out_png.unlink()
    generated.rename(out_png)


def build_all_docx(docx_dir: Path, items_by_code: dict[str, dict[int, dict]]) -> Path:
    document = setup_document()
    for index, page in enumerate(PAGES):
        if index > 0:
            document.add_page_break()
        build_page(document, page, items_by_code)
    out = docx_dir / "autismdev_training_effect_templates.docx"
    document.save(out)
    return out


def build_single_page_docx(docx_dir: Path, page: tuple[int, str, int, int], items_by_code: dict[str, dict[int, dict]]) -> Path:
    document = setup_document()
    build_page(document, page, items_by_code)
    page_no = page[0]
    out = docx_dir / f"page_{page_no:02d}.docx"
    document.save(out)
    return out


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx-dir", type=Path, default=DEFAULT_DOCX_DIR)
    parser.add_argument("--asset-dir", type=Path, default=ASSET_DIR)
    parser.add_argument("--skip-render", action="store_true")
    args = parser.parse_args()

    if not args.skip_render and shutil.which("qlmanage") is None:
        print("qlmanage is required to render DOCX thumbnails", file=sys.stderr)
        return 1

    args.docx_dir.mkdir(parents=True, exist_ok=True)
    args.asset_dir.mkdir(parents=True, exist_ok=True)
    items_by_code = load_items()
    all_docx = build_all_docx(args.docx_dir, items_by_code)
    print(f"wrote {all_docx}")
    for page in PAGES:
        page_no = page[0]
        docx = build_single_page_docx(args.docx_dir, page, items_by_code)
        print(f"wrote {docx}")
        if not args.skip_render:
            png = args.asset_dir / f"page_{page_no:02d}.png"
            draw_training_table_png(png, page_no, page[1], page[2], page[3], items_by_code[page[1]])
            print(f"rendered {png}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
